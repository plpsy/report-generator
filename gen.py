
import re
from docx import Document
from docx import text
import os

_item_col = 1
_option_col = 2
_expect_col = 3
_result_col = 4
_conclusion_col = 5

def get_output():
    with open('output.txt', 'r', encoding='utf-8') as f:
        lines = f.readlines()
        return [x.lower() for x in lines]
    return []

def target_exist(lines=[], key_str="", id_str="", target_str=""):
    match_str = key_str + id_str + target_str
    match_str = match_str.lower()
    for line in lines:
        if match_str in line:
            return True
    return False

def get_conclu_parag(row):
    parags = row.cells[_conclusion_col].paragraphs
    pass_parag = parags[0]
    fail_parag = parags[0]
    for p in parags:
        if("□合格" in p.text):
            pass_parag = p
        elif("□不合格" in p.text):
            fail_parag = p

    return pass_parag, fail_parag

def get_result_run(row, parag_idx):
    runs = row.cells[_result_col].paragraphs[parag_idx].runs
    result_run = runs[0]
    for run in runs:
        if(run.underline):
            result_run = run
            break
    return result_run

wildcards = ["xxxx", "yyyy", "zzzz"]

def is_wildcards(target_str):
    for w in wildcards:
        if(w in target_str):
            return True, target_str.split(w)
    return False, []

def infer_result(lines, row, target_str):
    is_wild, new_targets = is_wildcards(target_str)
    if(is_wild):
        left, right = new_targets[0].strip(), new_targets[1].strip()
        for line in lines:
            if left in line and right in line:
                result_str = line.split(left)[1]
                if(right == ''):
                    return True, f" {result_str.strip()} "
                else:
                    return True, f" {result_str.split(right)[0].strip()} "
        return False, " xxxx "
    else:
        for line in lines:
            if target_str in line:
                if "times" in target_str or "recv" in target_str:
                    count = re.findall(r'\d+', line)[-1]
                    return True, f" {count} "
                else:
                    return True, " Ok "
        else:
            return False, " Error "


def generate_row(lines, row):
    pass_parag, fail_parag = get_conclu_parag(row)

    all_success = True
    for j, p in enumerate(row.cells[_expect_col].paragraphs):
        result_run = get_result_run(row, j)
        ok, result_run.text = infer_result(lines, row, p.text.lower())
        if(not ok):
            all_success = False

    if(all_success):
        pass_parag.text="■合格"
    else:
        fail_parag.text="■不合格"

def save_to_docx():
    document = Document(docx=os.path.join(os.getcwd(), 'temple.docx'))    

    # for l, line in enumerate(get_output()):
    #     print(l)
    #     print(line)

    # paragraphs = document.paragraphs
    # for p in paragraphs:
    #     print(p.text)■

    for i, row in enumerate(document.tables[3].rows):
        for j, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                print(i,j)
                print(p.text)



def generate_table(document, lines, tabidx):
    for i, row in enumerate(document.tables[tabidx].rows):
        if(i == 0):
            continue
        option = row.cells[_option_col].paragraphs[0].text
        if option != "通过串口打印窗口读取" and option != "通过串口打印窗口观察":
            continue
        generate_row(lines, row)

def generate_docx():
    document = Document(docx=os.path.join(os.getcwd(), 'temple.docx'))
    lines = get_output()
    generate_table(document, lines, 2)
    generate_table(document, lines, 3)
    generate_table(document, lines, 4)
  

    document.save("test.docx")

def main():
    generate_docx()


main()